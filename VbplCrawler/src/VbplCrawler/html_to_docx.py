from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "li", "tr"}
SKIP_TAGS = {"script", "style", "head", "meta", "link"}


def _alignment_from_style(style: str) -> Optional[WD_ALIGN_PARAGRAPH]:
    normalized = style.replace(" ", "").lower()
    if "text-align:center" in normalized:
        return WD_ALIGN_PARAGRAPH.CENTER
    if "text-align:right" in normalized:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if "text-align:justify" in normalized:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return None


def _is_bold(node: Tag) -> bool:
    style = (node.get("style") or "").lower()
    if "font-weight:bold" in style.replace(" ", ""):
        return True
    return node.name in {"strong", "b"}


def _is_italic(node: Tag) -> bool:
    style = (node.get("style") or "").lower()
    if "font-style:italic" in style.replace(" ", ""):
        return True
    return node.name in {"em", "i"}


def _append_runs(paragraph, node, *, bold: bool = False, italic: bool = False) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        return

    if not isinstance(node, Tag):
        return

    if node.name in SKIP_TAGS:
        return

    if node.name == "br":
        paragraph.add_run().add_break()
        return

    child_bold = bold or _is_bold(node)
    child_italic = italic or _is_italic(node)

    for child in node.children:
        _append_runs(paragraph, child, bold=child_bold, italic=child_italic)


def _add_block(document: Document, node: Tag) -> None:
    text = node.get_text(" ", strip=True)
    if not text and node.name != "tr":
        return

    if node.name == "tr":
        cells = node.find_all(["td", "th"], recursive=False)
        if not cells:
            return
        table = document.add_table(rows=1, cols=len(cells))
        table.style = "Table Grid"
        for index, cell in enumerate(cells):
            paragraph = table.rows[0].cells[index].paragraphs[0]
            for child in cell.children:
                _append_runs(paragraph, child)
        return

    paragraph = document.add_paragraph()
    alignment = _alignment_from_style(node.get("style") or "")
    if alignment is not None:
        paragraph.alignment = alignment

    for child in node.children:
        _append_runs(paragraph, child)


def _iter_blocks(root: Tag):
    for node in root.descendants:
        if not isinstance(node, Tag):
            continue
        if node.name not in BLOCK_TAGS:
            continue
        if node.name == "tr" and node.find_parent("table") is not root.find_parent("table"):
            parent_table = node.find_parent("table")
            if parent_table and parent_table != root and parent_table.find_parent("table"):
                continue
        if node.name in {"p", "div", "h1", "h2", "h3", "h4", "li"}:
            if node.find_parent(BLOCK_TAGS - {node.name}):
                continue
        yield node


def _sanitize_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag_name in SKIP_TAGS:
        for node in soup.find_all(tag_name):
            node.decompose()
    for anchor in soup.find_all("a"):
        if not anchor.get("href"):
            anchor.unwrap()
    for image in soup.find_all("img"):
        image.decompose()
    body = soup.body or soup
    return str(body)


def html_to_docx(html: str, output_path: str | Path, *, title: Optional[str] = None) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(13)

    if title:
        heading = document.add_paragraph(title)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.bold = True
            run.font.size = Pt(14)

    soup = BeautifulSoup(_sanitize_html(html), "html.parser")
    root = soup.body or soup

    tables = set(root.find_all("table"))
    for node in _iter_blocks(root):
        if node.name == "tr":
            table = node.find_parent("table")
            if table is not None and table not in tables:
                continue
            if table is not None:
                tables.discard(table)
        _add_block(document, node)

    document.save(output)
    return output


def safe_filename(value: str, *, max_length: int = 120) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\n\r\t]+', "_", value).strip(" .")
    cleaned = re.sub(r"_+", "_", cleaned)
    if not cleaned:
        cleaned = "document"
    return cleaned[:max_length]
